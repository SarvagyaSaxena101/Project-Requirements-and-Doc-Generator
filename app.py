
import streamlit as st
import groq
from dotenv import load_dotenv
import os
import re
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
import io

def create_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    for line in text.split('\n'):
        if line.startswith('### '):
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, line[4:], 0, 1)
            pdf.set_font('Arial', '', 12)
        elif line.startswith('## '):
            pdf.set_font('Arial', 'B', 16)
            pdf.cell(0, 10, line[3:], 0, 1)
            pdf.set_font('Arial', '', 12)
        elif line.startswith('# '):
            pdf.set_font('Arial', 'B', 18)
            pdf.cell(0, 10, line[2:], 0, 1)
            pdf.set_font('Arial', '', 12)
        elif line.startswith('```'):
            pdf.set_font('Courier', '', 10)
            # Find the end of the code block
            code_block = line
            while '```' not in line:
                code_block += '\n' + line
            
            pdf.multi_cell(0, 5, code_block.replace('```', ''))
            pdf.set_font('Arial', '', 12)
        else:
            pdf.multi_cell(0, 10, line)
            
    pdf_output = pdf.output(dest='S').encode('latin1')
    return pdf_output

def create_docx(text):
    document = Document()
    
    # Ensure styles exist
    styles = document.styles
    if 'Heading 1' not in styles:
        styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH).base_style = styles['Normal']
    if 'Heading 2' not in styles:
        styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH).base_style = styles['Normal']
    if 'Heading 3' not in styles:
        styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH).base_style = styles['Normal']
    if 'Code' not in styles:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Courier New'

    for line in text.split('\n'):
        if line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('```'):
            p = document.add_paragraph(style='Code')
            p.add_run(line.replace('```', ''))
        else:
            document.add_paragraph(line)

    doc_output = io.BytesIO()
    document.save(doc_output)
    doc_output.seek(0)
    return doc_output.getvalue()

def parse_spec_from_response(response_text):
    sections = {
        "Functional Requirements": "",
        "Non-functional Requirements": "",
        "API Endpoints": "",
        "Database Schema": "",
        "Edge Cases": "",
        "Clarifying Questions": ""
    }
    
    # Split by the main sections first
    parts = re.split(r'### (Functional Requirements|Non-functional Requirements|API Endpoints|Database Schema|Edge Cases|Clarifying Questions)\s*\n', response_text)
    
    for i in range(1, len(parts), 2):
        header = parts[i].strip()
        content = parts[i+1].strip()
        if header in sections:
            sections[header] = content

    return sections

def main():
    load_dotenv()
    st.title("AutoSpec AI – Natural Language → Technical Specification Generator")

    user_input = st.text_area("Describe your project idea:", "I want a food delivery app for hostels with cash-on-delivery and time slots.")

    if st.button("Generate Specification"):
        groq_api_key = os.environ.get("GROQ_API_KEY")
        if not groq_api_key or groq_api_key == "YOUR_GROQ_API_KEY":
            st.error("Please add your GROQ_API_KEY to the .env file.")
            return
        generate_specification(user_input)

def generate_specification(user_input):
    try:
        client = groq.Groq(api_key=os.environ.get("GROQ_API_KEY"))
        
        prompt = f"""
        Generate a comprehensive and detailed technical specification for the following project idea: "{user_input}"

        For each section, provide in-depth analysis, examples, and detailed descriptions. Use markdown for formatting. 
        Where applicable, include ASCII diagrams to illustrate concepts like system architecture or database relationships.

        The specification must include the following sections, with each section starting with a '###' header:
        - ### Functional Requirements (Provide a detailed list with sub-headings for each feature)
        - ### Non-functional Requirements (Cover aspects like performance, security, scalability, etc. in detail)
        - ### API Endpoints (List all endpoints with request/response examples for each)
        - ### Database Schema (Provide a detailed schema with table names, columns, data types, and relationships. Include an ASCII diagram of the schema.)
        - ### Edge Cases (List potential edge cases for each major feature)
        - ### Clarifying Questions (Ask detailed questions to resolve ambiguities in the project idea)
        """

        chat_completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
        )
        response_text = chat_completion.choices[0].message.content
        parsed_spec = parse_spec_from_response(response_text)

        st.header("Generated Specification")

        # Functional Requirements
        with st.expander("Functional Requirements"):
            fr_text = st.text_area("Edit Functional Requirements", parsed_spec.get("Functional Requirements", ""), key="fr_edit")
            st.markdown(fr_text)
            st.download_button("Download as PDF", create_pdf(fr_text), "functional_requirements.pdf", "application/pdf", key="fr_pdf")
            st.download_button("Download as DOCX", create_docx(fr_text), "functional_requirements.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="fr_docx")

        # Other sections follow the same pattern...
        with st.expander("Non-functional Requirements"):
            nfr_text = st.text_area("Edit Non-functional Requirements", parsed_spec.get("Non-functional Requirements", ""), key="nfr_edit")
            st.markdown(nfr_text)
            st.download_button("Download as PDF", create_pdf(nfr_text), "non_functional_requirements.pdf", "application/pdf", key="nfr_pdf")
            st.download_button("Download as DOCX", create_docx(nfr_text), "non_functional_requirements.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="nfr_docx")

        with st.expander("API Endpoints"):
            api_text = st.text_area("Edit API Endpoints", parsed_spec.get("API Endpoints", ""), key="api_edit")
            st.markdown(api_text)
            st.download_button("Download as PDF", create_pdf(api_text), "api_endpoints.pdf", "application/pdf", key="api_pdf")
            st.download_button("Download as DOCX", create_docx(api_text), "api_endpoints.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="api_docx")

        with st.expander("Database Schema"):
            db_text = st.text_area("Edit Database Schema", parsed_spec.get("Database Schema", ""), key="db_edit")
            st.markdown(db_text)
            st.download_button("Download as PDF", create_pdf(db_text), "database_schema.pdf", "application/pdf", key="db_pdf")
            st.download_button("Download as DOCX", create_docx(db_text), "database_schema.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="db_docx")

        with st.expander("Edge Cases"):
            ec_text = st.text_area("Edit Edge Cases", parsed_spec.get("Edge Cases", ""), key="ec_edit")
            st.markdown(ec_text)
            st.download_button("Download as PDF", create_pdf(ec_text), "edge_cases.pdf", "application/pdf", key="ec_pdf")
            st.download_button("Download as DOCX", create_docx(ec_text), "edge_cases.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="ec_docx")

        with st.expander("Clarifying Questions"):
            cq_text = st.text_area("Edit Clarifying Questions", parsed_spec.get("Clarifying Questions", ""), key="cq_edit")
            st.markdown(cq_text)
            st.download_button("Download as PDF", create_pdf(cq_text), "clarifying_questions.pdf", "application/pdf", key="cq_pdf")
            st.download_button("Download as DOCX", create_docx(cq_text), "clarifying_questions.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="cq_docx")

    except Exception as e:
        st.error(f"An error occurred: {e}")

if __name__ == "__main__":
    main()
